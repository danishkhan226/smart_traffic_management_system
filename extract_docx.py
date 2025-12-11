from docx import Document
import sys

try:
    # Open the document
    doc = Document('conference-Papertemplate.docx')
    
    print("=" * 80)
    print("DOCUMENT CONTENT:")
    print("=" * 80)
    
    # Extract all paragraphs
    for i, para in enumerate(doc.paragraphs, 1):
        if para.text.strip():  # Only print non-empty paragraphs
            print(f"\n[Paragraph {i}]")
            print(para.text)
    
    # Extract tables if any
    if doc.tables:
        print("\n" + "=" * 80)
        print("TABLES:")
        print("=" * 80)
        for i, table in enumerate(doc.tables, 1):
            print(f"\n[Table {i}]")
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                print(" | ".join(row_text))
    
except Exception as e:
    print(f"Error: {e}")
    import traceback
    traceback.print_exc()
